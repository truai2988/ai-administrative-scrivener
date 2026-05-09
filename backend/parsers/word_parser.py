"""
Word パーサー

python-docx を用いてWordファイル（.docx）からテキストと表データを抽出する。
段落、表、ヘッダー/フッターの情報を保持する。
"""

import io
import logging

import docx
import docx.oxml.ns

logger = logging.getLogger(__name__)

# 段落数の上限（OOMおよびタイムアウト防止）
_MAX_PARAGRAPHS = 3000
# 表の行数上限
_MAX_TABLE_ROWS = 500


def extract_word_text(file_bytes: bytes, filename: str) -> str:
    """
    Wordファイルのバイト列を受け取り、構造化テキストに変換する。

    - 段落テキストをそのまま出力
    - 表データを行ごとに「|」区切りで出力
    - ヘッダー/フッターも抽出
    - 段落数・表の行数に上限を設定してOOM/タイムアウトを防止
    """
    try:
        document = docx.Document(io.BytesIO(file_bytes))
    except Exception as e:
        logger.error(f"python-docx でのファイル読み込みに失敗: {filename} - {e}")
        return f"=== ファイル: {filename} (読み込みエラー) ===\n"

    sections: list[str] = [f"=== ファイル: {filename} ===\n"]

    # ヘッダー抽出
    for section in document.sections:
        header = section.header
        if header and header.paragraphs:
            header_text = " ".join(p.text.strip() for p in header.paragraphs if p.text.strip())
            if header_text:
                sections.append(f"[ヘッダー] {header_text}")

    # 本文の段落と表を出現順に処理
    para_count = 0
    table_row_total = 0
    truncated = False

    for element in document.element.body:
        tag = element.tag.split("}")[-1]  # 名前空間を除去

        if tag == "p":
            para_count += 1
            if para_count > _MAX_PARAGRAPHS:
                if not truncated:
                    sections.append(f"\n... (段落数上限 {_MAX_PARAGRAPHS} に到達。以降省略) ...")
                    truncated = True
                continue

            # 段落
            para = docx.oxml.ns.qn("w:p")
            if element.tag == para or tag == "p":
                text = element.text or ""
                # python-docx の Paragraph オブジェクトから取得
                for para_obj in document.paragraphs:
                    if para_obj._element is element:
                        text = para_obj.text.strip()
                        break
                if text:
                    sections.append(f"  {text}")

        elif tag == "tbl":
            # 表
            sections.append("\n  [表]")
            for table in document.tables:
                if table._element is element:
                    row_count = 0
                    for row in table.rows:
                        row_count += 1
                        table_row_total += 1
                        if row_count > _MAX_TABLE_ROWS:
                            sections.append(f"    ... (表の行数上限 {_MAX_TABLE_ROWS} に到達。以降省略)")
                            break
                        cells = [cell.text.strip() for cell in row.cells]
                        sections.append(f"    {' | '.join(cells)}")
                    sections.append("")
                    break

    # フッター抽出
    for section in document.sections:
        footer = section.footer
        if footer and footer.paragraphs:
            footer_text = " ".join(p.text.strip() for p in footer.paragraphs if p.text.strip())
            if footer_text:
                sections.append(f"[フッター] {footer_text}")

    result_text = "\n".join(sections)
    logger.info(
        f"Word解析完了: {filename} "
        f"(段落数: {para_count}, 表行数合計: {table_row_total}, "
        f"テキスト長: {len(result_text)} 文字)"
    )
    return result_text
